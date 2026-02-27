import os
import pandas as pd
from docx import Document
from pypdf import PdfReader, PdfWriter
from package.core_utils.log_manager import LogManager
import platform
import subprocess

logger = LogManager.get_logger(__name__)

def open_in_native_app(file_path):
    """Opens a file in its default application."""
    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        return False

    try:
        if platform.system() == 'Windows':
            os.startfile(file_path)
        elif platform.system() == 'Darwin':
            subprocess.run(['open', file_path], shell=False)
        else:
            subprocess.run(['xdg-open', file_path], shell=False)
        return True
    except Exception as e:
        logger.error(f"Error opening file {file_path}: {e}")
        return False

class OfficeAutomator:
    @staticmethod
    def create_excel_report(data, output_path, sheet_name="Report"):
        """Creates an Excel report from a list of dicts or a DataFrame."""
        try:
            if isinstance(data, pd.DataFrame):
                df = data
            else:
                df = pd.DataFrame(data)

            df.to_excel(output_path, index=False, sheet_name=sheet_name)
            logger.info(f"Excel report created at {output_path}")
            return True
        except Exception as e:
            logger.error(f"Failed to create Excel report: {e}")
            return False

    @staticmethod
    def create_word_document(content, output_path, title=None):
        """Creates a Word document with the given content."""
        try:
            doc = Document()
            if title:
                doc.add_heading(title, 0)

            if isinstance(content, list):
                for item in content:
                    doc.add_paragraph(str(item))
            else:
                doc.add_paragraph(str(content))

            doc.save(output_path)
            logger.info(f"Word document saved at {output_path}")
            return True
        except Exception as e:
            logger.error(f"Failed to create Word document: {e}")
            return False

    @staticmethod
    def fill_pdf_fields(input_pdf, output_pdf, field_data):
        """Fills fields in a PDF form."""
        # Note: filling PDF forms programmatically can be complex depending on the PDF structure.
        # This is a basic implementation using pypdf.
        try:
            reader = PdfReader(input_pdf)
            writer = PdfWriter()

            for page in reader.pages:
                writer.add_page(page)

            writer.update_page_form_field_values(writer.pages[0], field_data)

            with open(output_pdf, "wb") as output_stream:
                writer.write(output_stream)

            logger.info(f"Filled PDF saved at {output_pdf}")
            return True
        except Exception as e:
            logger.error(f"Failed to fill PDF: {e}")
            return False

automator = OfficeAutomator()
