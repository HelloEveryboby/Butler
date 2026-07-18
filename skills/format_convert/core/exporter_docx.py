import re
import os
import logging

logger = logging.getLogger("DocxExporter")

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def _set_cell_background(cell, color_hex):
    """Sets background color of a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    tc_pr.append(parse_xml(shd_xml))


def _set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets padding for a table cell in twentieths of a point (dxa)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def markdown_to_docx(md_text, output_path, opts=None):
    """
    Converts Markdown to DOCX with high-fidelity Apple-style minimal aesthetic.
    Supports headings, lists, bold/italic, code blocks, quote blocks, tables, and rules.
    """
    if not DOCX_AVAILABLE:
        # Graceful fallback: write basic structured text if docx package is missing
        logger.warning("python-docx is not installed. Saving formatted TXT file as a fallback.")
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("=== Fallback Document (python-docx not installed) ===\n\n")
            f.write(md_text)
        return output_path

    opts = opts or {}
    theme = opts.get("theme", "apple-light")
    with_water = opts.get("with_water", False)

    doc = Document()

    # Define Apple aesthetic color palettes
    if theme == "dark":
        text_color = RGBColor(0xEE, 0xEE, 0xEE)
        h1_color = RGBColor(0xFF, 0xFF, 0xFF)
        quote_bg = "1C1C1E"
        code_bg = "2C2C2E"
        bg_color_hex = "121212"
        # Since standard Word pages are white, we maintain dark elements on light background
        # or dark background styling if preferred. To keep standard readability, we use dark text design.
        text_color = RGBColor(0x1D, 0x1D, 0x1F)
        h1_color = RGBColor(0x1D, 0x1D, 0x1F)
        quote_bg = "F4F4F4"
        code_bg = "F5F5F7"
    else:
        text_color = RGBColor(0x1D, 0x1D, 0x1F)
        h1_color = RGBColor(0x1D, 0x1D, 0x1F)
        quote_bg = "F4F4F4"
        code_bg = "F5F5F7"

    # Document margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Helper to style a run with SF Pro / PingFang style font
    def apply_font_run(run, size_pt, bold=False, italic=False, color=text_color):
        run.font.name = 'PingFang SC'
        run.font.size = Pt(size_pt)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = color
        # Set East Asia font specifically
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), 'PingFang SC')
        rFonts.set(qn('w:ascii'), 'SF Pro Text')
        rFonts.set(qn('w:hAnsi'), 'SF Pro Text')
        rPr.append(rFonts)

    # Process block markdown
    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^(?:-{3,}|\*{3,}|\_{3,})$', stripped):
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.space_before = Pt(12)
            p_format.space_after = Pt(12)
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'E5E5EA')
            pBdr.append(bottom)
            p._p.get_or_add_pPr().append(pBdr)
            i += 1
            continue

        # Fenced Code Block
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1 # skip ending ```

            # Render code block inside a styled 1x1 table (like a box)
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell = table.cell(0, 0)
            _set_cell_background(cell, code_bg)
            _set_cell_margins(cell, top=140, bottom=140, left=200, right=200)

            # Set a light gray border for code block
            tcBorders = OxmlElement('w:tcBorders')
            for b_name in ['top', 'left', 'bottom', 'right']:
                b_el = OxmlElement(f'w:{b_name}')
                b_el.set(qn('w:val'), 'single')
                b_el.set(qn('w:sz'), '4')
                b_el.set(qn('w:color'), 'E5E5EA')
                tcBorders.append(b_el)
            cell._tc.get_or_add_tcPr().append(tcBorders)

            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run("\n".join(code_lines))
            apply_font_run(run, 9.5, color=RGBColor(0x3A, 0x3A, 0x3C))
            run.font.name = 'Courier New'
            continue

        # Headings
        heading_match = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            title_text = heading_match.group(2)
            p = doc.add_paragraph()
            p_format = p.paragraph_format

            if level == 1:
                p_format.space_before = Pt(18)
                p_format.space_after = Pt(8)
                run = p.add_run(title_text)
                apply_font_run(run, 20, bold=True, color=h1_color)
            elif level == 2:
                p_format.space_before = Pt(14)
                p_format.space_after = Pt(6)
                run = p.add_run(title_text)
                apply_font_run(run, 16, bold=True, color=h1_color)
            elif level == 3:
                p_format.space_before = Pt(12)
                p_format.space_after = Pt(4)
                run = p.add_run(title_text)
                apply_font_run(run, 13, bold=True, color=h1_color)
            else:
                p_format.space_before = Pt(10)
                p_format.space_after = Pt(4)
                run = p.add_run(title_text)
                apply_font_run(run, 11, bold=True, color=h1_color)

            i += 1
            continue

        # Blockquote (starting with '>')
        if stripped.startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(re.sub(r'^>\s*', '', lines[i].strip()))
                i += 1

            # Quote block layout: 1x1 table with thick left border, gray background
            table = doc.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            _set_cell_background(cell, quote_bg)
            _set_cell_margins(cell, top=100, bottom=100, left=180, right=150)

            # Thick left border, none for others
            tcBorders = OxmlElement('w:tcBorders')
            left_b = OxmlElement('w:left')
            left_b.set(qn('w:val'), 'single')
            left_b.set(qn('w:sz'), '24') # 3pt
            left_b.set(qn('w:color'), '007AFF') # Apple Blue accent
            tcBorders.append(left_b)
            for b_name in ['top', 'bottom', 'right']:
                b_el = OxmlElement(f'w:{b_name}')
                b_el.set(qn('w:val'), 'none')
                tcBorders.append(b_el)
            cell._tc.get_or_add_tcPr().append(tcBorders)

            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run("\n".join(quote_lines))
            apply_font_run(run, 10, italic=True, color=RGBColor(0x55, 0x55, 0x55))
            continue

        # Tables
        if stripped.startswith("|") and i < len(lines):
            # Parse table rows
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_stripped = lines[i].strip()
                # Split and filter empty cells at ends
                cells = [c.strip() for c in row_stripped.split("|")[1:-1]]
                table_rows.append(cells)
                i += 1

            if len(table_rows) > 0:
                # Filter separator row (e.g. |---|---|)
                clean_rows = []
                for idx, r in enumerate(table_rows):
                    if idx == 1 and all(re.match(r'^:?-+:?$', c) for c in r):
                        continue
                    clean_rows.append(r)

                if clean_rows:
                    cols_count = max(len(r) for r in clean_rows)
                    word_table = doc.add_table(rows=len(clean_rows), cols=cols_count)
                    word_table.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    for row_idx, r_cells in enumerate(clean_rows):
                        for col_idx, cell_val in enumerate(r_cells):
                            if col_idx < len(word_table.columns):
                                cell = word_table.cell(row_idx, col_idx)
                                cell.text = ""
                                _set_cell_margins(cell, top=80, bottom=80, left=100, right=100)

                                p = cell.paragraphs[0]
                                p.paragraph_format.space_before = Pt(0)
                                p.paragraph_format.space_after = Pt(0)

                                if row_idx == 0:
                                    # Header row
                                    _set_cell_background(cell, "F2F2F7")
                                    run = p.add_run(cell_val)
                                    apply_font_run(run, 10, bold=True, color=RGBColor(0x1D, 0x1D, 0x1F))
                                else:
                                    # Data row zebra striping
                                    if row_idx % 2 == 0:
                                        _set_cell_background(cell, "FAFAFC")
                                    run = p.add_run(cell_val)
                                    apply_font_run(run, 10, color=RGBColor(0x3A, 0x3A, 0x3C))

                                # Apply borders to all cells
                                tcBorders = OxmlElement('w:tcBorders')
                                for b_name in ['top', 'left', 'bottom', 'right']:
                                    b_el = OxmlElement(f'w:{b_name}')
                                    b_el.set(qn('w:val'), 'single')
                                    b_el.set(qn('w:sz'), '4')
                                    b_el.set(qn('w:color'), 'E5E5EA')
                                    tcBorders.append(b_el)
                                cell._tc.get_or_add_tcPr().append(tcBorders)
            continue

        # Lists (unordered/ordered)
        unordered_match = re.match(r'^([\-\*\+])\s+(.*)$', stripped)
        ordered_match = re.match(r'^(\d+)\.\s+(.*)$', stripped)

        if unordered_match or ordered_match:
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.space_before = Pt(0)
            p_format.space_after = Pt(3)
            p_format.left_indent = Inches(0.25)

            if unordered_match:
                # Add elegant list bullet
                run_bullet = p.add_run("•  ")
                apply_font_run(run_bullet, 11, bold=True, color=RGBColor(0x00, 0x7A, 0xFF))
                content = unordered_match.group(2)
            else:
                num_prefix = ordered_match.group(1)
                run_num = p.add_run(f"{num_prefix}.  ")
                apply_font_run(run_num, 11, bold=True, color=RGBColor(0x00, 0x7A, 0xFF))
                content = ordered_match.group(2)

            _add_formatted_inline(p, content, apply_font_run, text_color)
            i += 1
            continue

        # Normal Paragraph
        p = doc.add_paragraph()
        p_format = p.paragraph_format
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(8)
        p_format.line_spacing = 1.25

        _add_formatted_inline(p, stripped, apply_font_run, text_color)
        i += 1

    # Watermark Footer
    if with_water:
        section = doc.sections[0]
        footer = section.footer
        f_p = footer.paragraphs[0]
        f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = f_p.add_run("Generated by Butler Document Workspace | Secure Watermarked")
        apply_font_run(run, 8.5, color=RGBColor(0x8E, 0x8E, 0x93))

    doc.save(output_path)
    return output_path


def _add_formatted_inline(paragraph, text, font_setter, text_color):
    """Parses and adds inline Markdown formats (bold, italic, inline-code)."""
    # Regex for bold (**text** or __text__), italic (*text* or _text_), inline code (`code`)
    pattern = re.compile(r'(\*\*.*?\*\*|__.*?__|`.*?`|\*.*?\*|_.*?_)')
    parts = pattern.split(text)

    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            font_setter(run, 11, bold=True, color=text_color)
        elif part.startswith("__") and part.endswith("__"):
            run = paragraph.add_run(part[2:-2])
            font_setter(run, 11, bold=True, color=text_color)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            font_setter(run, 11, italic=True, color=text_color)
        elif part.startswith("_") and part.endswith("_"):
            run = paragraph.add_run(part[1:-1])
            font_setter(run, 11, italic=True, color=text_color)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            font_setter(run, 10, color=RGBColor(0xC9, 0x3B, 0x75)) # Classic pink code block
            run.font.name = 'Courier New'
        else:
            run = paragraph.add_run(part)
            font_setter(run, 11, color=text_color)
