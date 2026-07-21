import logging
import os
import re

logger = logging.getLogger("ReverseConverters")


def _try_markitdown(file_path_or_bytes, suffix=".docx"):
    """
    Attempts to convert various formats to Markdown using:
    1. Microsoft's official `markitdown` library
    2. Butler's internal replica `skills.markitdown.markitdown_app`
    Returns converted Markdown string or None.
    """
    # Create temp file if input is bytes
    temp_path = None
    if isinstance(file_path_or_bytes, bytes):
        import tempfile

        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(file_path_or_bytes)
            temp_path = tmp.name
        file_path = temp_path
    else:
        file_path = file_path_or_bytes

    try:
        # Strategy A: Microsoft's official markitdown package
        try:
            from markitdown import MarkItDown

            mid = MarkItDown()
            result = mid.convert(file_path)
            return result.text_content
        except ImportError:
            pass

        # Strategy B: Butler's custom integrated markitdown_app replica
        try:
            # Add project root to path for absolute imports if needed
            import sys

            project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
            if project_root not in sys.path:
                sys.path.insert(0, project_root)

            from skills.markitdown.markitdown_app import convert as butler_convert

            return butler_convert(file_path)
        except Exception as e:
            logger.info(f"Butler internal markitdown conversion skipped: {e}")

    except Exception as e:
        logger.warning(f"Markitdown strategy failed: {e}")
    finally:
        if temp_path and os.path.exists(temp_path):
            os.remove(temp_path)

    return None


def html_to_markdown(html_text):
    """
    Converts HTML back to clean Markdown format.
    """
    # Try markitdown first
    mid_res = _try_markitdown(html_text.encode("utf-8"), suffix=".html")
    if mid_res:
        return mid_res

    # Strategy C: markdownify
    try:
        from markdownify import markdownify as md

        return md(html_text)
    except ImportError:
        logger.info("markdownify not installed. Running high-strength HTML to Markdown fallback regex parser.")

        text = html_text

        # Strip head, scripts, styling metadata
        text = re.sub(r"<(style|script|head)\b[^>]*>([\s\S]*?)<\/\1>", "", text, flags=re.IGNORECASE)

        # Headings mapping
        text = re.sub(r"<h1\b[^>]*>(.*?)</h1>", r"# \1\n\n", text, flags=re.IGNORECASE)
        text = re.sub(r"<h2\b[^>]*>(.*?)</h2>", r"## \1\n\n", text, flags=re.IGNORECASE)
        text = re.sub(r"<h3\b[^>]*>(.*?)</h3>", r"### \1\n\n", text, flags=re.IGNORECASE)

        # Lists mapping
        text = re.sub(r"<li\b[^>]*>(.*?)</li>", r"- \1\n", text, flags=re.IGNORECASE)
        text = re.sub(r"</?(ul|ol)>", r"\n", text, flags=re.IGNORECASE)

        # Code blocks mapping
        text = re.sub(
            r"<pre\b[^>]*><code\b[^>]*>([\s\S]*?)</code></pre>", r"```\n\1\n```\n\n", text, flags=re.IGNORECASE
        )
        text = re.sub(r"<code\b[^>]*>(.*?)</code>", r"`\1`", text, flags=re.IGNORECASE)

        # Bold & Italic mapping
        text = re.sub(r"<(strong|b)\b[^>]*>(.*?)</\1>", r"**\2**", text, flags=re.IGNORECASE)
        text = re.sub(r"<(em|i)\b[^>]*>(.*?)</\1>", r"*\2*", text, flags=re.IGNORECASE)

        # Paragraphs & line breaks
        text = re.sub(r"<p\b[^>]*>(.*?)</p>", r"\1\n\n", text, flags=re.IGNORECASE)
        text = re.sub(r"<br\s*/?>", r"\n", text, flags=re.IGNORECASE)

        # Strip remaining HTML tags
        text = re.sub(r"<[^>]+>", "", text)

        # Decode standard entities
        text = text.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">").replace("&quot;", '"')

        return text.strip()


def docx_to_markdown(docx_path_or_bytes):
    """
    Parses a DOCX document back into structured Markdown notation.
    """
    # Try markitdown first
    mid_res = _try_markitdown(docx_path_or_bytes, suffix=".docx")
    if mid_res:
        return mid_res

    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for reverse DOCX -> MD conversion.")

    # Load from path or bytes stream
    if isinstance(docx_path_or_bytes, bytes):
        import io

        doc = Document(io.BytesIO(docx_path_or_bytes))
    else:
        doc = Document(docx_path_or_bytes)

    md_lines = []
    for p in doc.paragraphs:
        text = ""
        # Inspect runs for inline formatting (bold, italic)
        for run in p.runs:
            r_text = run.text
            if not r_text:
                continue
            if run.bold:
                r_text = f"**{r_text}**"
            if run.italic:
                r_text = f"*{r_text}*"
            text += r_text

        # Check paragraph style / structure
        style_name = p.style.name.lower()
        if "heading 1" in style_name:
            md_lines.append(f"# {text}\n")
        elif "heading 2" in style_name:
            md_lines.append(f"## {text}\n")
        elif "heading 3" in style_name:
            md_lines.append(f"### {text}\n")
        elif "list bullet" in style_name or p.paragraph_format.left_indent:
            md_lines.append(f"- {text}")
        else:
            if text.strip():
                md_lines.append(text + "\n")

    # Parse tables as markdown tables if present
    for table in doc.tables:
        table_md = []
        for r_idx, row in enumerate(table.rows):
            row_cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
            table_md.append("| " + " | ".join(row_cells) + " |")
            if r_idx == 0:
                table_md.append("| " + " | ".join("---" for _ in row.cells) + " |")
        md_lines.append("\n" + "\n".join(table_md) + "\n")

    return "\n".join(md_lines).strip()


def pdf_to_markdown(pdf_path_or_bytes):
    """
    Parses structural text from PDF pages and returns standard markdown string.
    """
    # Try markitdown first
    mid_res = _try_markitdown(pdf_path_or_bytes, suffix=".pdf")
    if mid_res:
        return mid_res

    # Try pypdf
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("pypdf is required for reverse PDF -> MD conversion.")

    import io

    if isinstance(pdf_path_or_bytes, bytes):
        reader = PdfReader(io.BytesIO(pdf_path_or_bytes))
    else:
        reader = PdfReader(pdf_path_or_bytes)

    extracted_lines = []
    for page in reader.pages:
        txt = page.extract_text()
        if txt:
            extracted_lines.append(txt)

    content = "\n\n".join(extracted_lines)
    return content.strip()
